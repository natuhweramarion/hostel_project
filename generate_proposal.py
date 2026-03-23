"""
Project Proposal Generator — Online Hostel Information Management System
Natuhwera Marion | Cavendish University Uganda
Run once to produce the Word document, then delete this script if desired.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── helpers ──────────────────────────────────────────────────────────────────

NAVY   = RGBColor(0x0D, 0x1F, 0x3C)
GOLD   = RGBColor(0xF0, 0xA5, 0x00)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
GRAY   = RGBColor(0x64, 0x74, 0x8B)
LGRAY  = RGBColor(0xF1, 0xF4, 0xF8)

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color=NAVY, size=16, bold=True, center=False):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold      = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p

def add_body(doc, text, indent=False, size=11, italic=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic    = italic
    return p

def add_bullet(doc, text, size=11):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

def add_numbered(doc, text, size=11):
    p = doc.add_paragraph(style='List Number')
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

def section_rule(doc, color_hex="0D1F3C"):
    """Thin coloured horizontal rule implemented as a 1-row, 1-col table."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    set_cell_bg(cell, color_hex)
    cell.height = Cm(0.07)
    cell._tc.get_or_add_tcPr()
    doc.add_paragraph()   # spacing after rule

# ── document setup ────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# Default paragraph style
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ── TITLE PAGE ────────────────────────────────────────────────────────────────

doc.add_paragraph()
doc.add_paragraph()

add_heading(doc, "CAVENDISH UNIVERSITY UGANDA", level=1,
            color=NAVY, size=20, center=True)
add_heading(doc, "Department of Science & Information Technology",
            level=2, color=GRAY, size=13, bold=False, center=True)

doc.add_paragraph()
section_rule(doc, "F0A500")   # gold rule

doc.add_paragraph()
add_heading(doc, "PROJECT PROPOSAL", color=GOLD, size=22, center=True)
doc.add_paragraph()
add_heading(doc,
    "Online Hostel Information Management System",
    color=NAVY, size=17, center=True)
doc.add_paragraph()
section_rule(doc, "0D1F3C")   # navy rule
doc.add_paragraph()

# Meta table
meta = doc.add_table(rows=6, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
meta.style = 'Table Grid'
data = [
    ("Student Name",    "Natuhwera Marion"),
    ("Student ID",      "To be filled"),
    ("Programme",       "Bachelor of Science in Information Technology"),
    ("Supervisor",      "To be assigned"),
    ("Academic Year",   "2025 / 2026"),
    ("Date",            datetime.date.today().strftime("%B %d, %Y")),
]
for i, (label, value) in enumerate(data):
    row = meta.rows[i]
    set_cell_bg(row.cells[0], "0D1F3C")
    lrun = row.cells[0].paragraphs[0].add_run(label)
    lrun.bold            = True
    lrun.font.color.rgb  = WHITE
    lrun.font.size       = Pt(10.5)
    vrun = row.cells[1].paragraphs[0].add_run(value)
    vrun.font.size       = Pt(10.5)

doc.add_page_break()

# ── 1. INTRODUCTION ───────────────────────────────────────────────────────────

add_heading(doc, "1. INTRODUCTION", size=14)
section_rule(doc, "F0A500")

add_body(doc,
    "The management of student accommodation in universities and higher institutions of learning "
    "has traditionally been a manual, paper-based process. This approach is prone to errors, "
    "inefficiencies, and difficulty in real-time tracking of room allocations, payments, and "
    "occupancy rates. Cavendish University Uganda, like many academic institutions, faces these "
    "challenges in administering its hostel facilities."
)
doc.add_paragraph()
add_body(doc,
    "This project proposes the design and implementation of an Online Hostel Information "
    "Management System (HMS) — a web-based platform built with Django (Python) that digitises "
    "the end-to-end workflow of hostel administration, from room allocation and payment "
    "verification to real-time occupancy reporting."
)

doc.add_paragraph()

# ── 2. PROBLEM STATEMENT ──────────────────────────────────────────────────────

add_heading(doc, "2. PROBLEM STATEMENT", size=14)
section_rule(doc, "F0A500")

add_body(doc,
    "Cavendish University Uganda currently manages hostel accommodation using manual registers "
    "and spreadsheets. This creates several operational problems:"
)
problems = [
    "Difficulty tracking which rooms are occupied, available, or fully booked in real time.",
    "Manual allocation of students to rooms leads to errors such as double-booking.",
    "Payment records are maintained on paper, making verification slow and prone to loss.",
    "There is no centralised system for generating occupancy or payment reports.",
    "Students have no self-service portal to view their allocation status or submit payments.",
    "Hostel wardens cannot quickly identify capacity issues or flag pending verifications.",
]
for p in problems:
    add_bullet(doc, p)

doc.add_paragraph()

# ── 3. OBJECTIVES ─────────────────────────────────────────────────────────────

add_heading(doc, "3. OBJECTIVES", size=14)
section_rule(doc, "F0A500")

add_heading(doc, "3.1 General Objective", size=12, color=NAVY, bold=True)
add_body(doc,
    "To design, develop, and deploy an Online Hostel Information Management System that "
    "automates student accommodation processes at Cavendish University Uganda."
)
doc.add_paragraph()

add_heading(doc, "3.2 Specific Objectives", size=12, color=NAVY, bold=True)
specific = [
    "To identify and document the requirements of the hostel management system through "
    "stakeholder analysis and user research.",
    "To design the system architecture, database schema, user interface, and role-based "
    "access model for the hostel management system.",
    "To implement the system using Django (Python), SQLite, Bootstrap 5, and Django's "
    "built-in authentication framework.",
    "To evaluate the system through functional testing, usability testing, and stakeholder "
    "review to ensure it meets defined requirements.",
]
for obj in specific:
    add_numbered(doc, obj)

doc.add_paragraph()

# ── 4. SCOPE ──────────────────────────────────────────────────────────────────

add_heading(doc, "4. SCOPE OF THE PROJECT", size=14)
section_rule(doc, "F0A500")

add_heading(doc, "4.1 In Scope", size=12, color=NAVY, bold=True)
in_scope = [
    "Student registration, login, and profile management.",
    "Hostel, block, and room inventory management (three-level hierarchy).",
    "Manual room allocation by hostel administrators with capacity and gender validation.",
    "Student-initiated payment submission with reference number and payment method.",
    "Payment verification and rejection by administrators.",
    "Role-based dashboards: Admin Dashboard and Student Dashboard.",
    "Occupancy, allocation, and payment reports with CSV export.",
    "Available rooms view for administrators.",
    "Secure session-based authentication with POST-only logout.",
]
for item in in_scope:
    add_bullet(doc, item)

doc.add_paragraph()
add_heading(doc, "4.2 Out of Scope", size=12, color=NAVY, bold=True)
out_scope = [
    "Automatic room allocation algorithms.",
    "Online payment gateway integration (e.g., Flutterwave, MTN Mobile Money).",
    "Email or SMS notification services.",
    "Mobile application (Android/iOS).",
    "Multi-institution or multi-tenancy support.",
    "REST API or third-party system integrations.",
]
for item in out_scope:
    add_bullet(doc, item)

doc.add_paragraph()

# ── 5. SIGNIFICANCE ───────────────────────────────────────────────────────────

add_heading(doc, "5. SIGNIFICANCE OF THE STUDY", size=14)
section_rule(doc, "F0A500")

significance = [
    ("To the University",
     "Provides a centralised, digital system that reduces administrative workload, eliminates "
     "manual record-keeping errors, and enables real-time visibility of hostel occupancy."),
    ("To Students",
     "Gives students a self-service portal to view their room allocation, track payment "
     "status, and access accommodation information without visiting the hostel office."),
    ("To Administrators",
     "Enables hostel wardens and finance staff to manage allocations, verify payments, and "
     "generate reports quickly and accurately."),
    ("To the Developer",
     "Provides practical experience in full-stack web development, software engineering "
     "principles, database design, and project management."),
    ("To Future Researchers",
     "Serves as a reference implementation for hostel or accommodation management systems "
     "in similar institutional contexts."),
]
for title, body in significance:
    p = doc.add_paragraph()
    r1 = p.add_run(title + ": ")
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(body)
    r2.font.size = Pt(11)

doc.add_paragraph()

# ── 6. METHODOLOGY ────────────────────────────────────────────────────────────

add_heading(doc, "6. METHODOLOGY", size=14)
section_rule(doc, "F0A500")

add_body(doc,
    "The project follows the Software Development Life Cycle (SDLC) using an iterative "
    "Agile-inspired approach, broken into the following phases:"
)
doc.add_paragraph()

phases = [
    ("Phase 1 — Requirements Gathering",
     "Conduct interviews and observations with hostel administrators and students at "
     "Cavendish University Uganda to identify functional and non-functional requirements. "
     "Document use cases, user stories, and system constraints."),
    ("Phase 2 — System Design",
     "Produce system architecture diagrams, Entity-Relationship (ER) diagrams, data flow "
     "diagrams (DFD), use case diagrams, and wireframes/mockups for all key screens. "
     "Define the role-based access control model."),
    ("Phase 3 — Implementation",
     "Develop the system using Django 5.2 (Python), SQLite for the database, Bootstrap 5.3 "
     "for the front end, and Bootstrap Icons for the UI. Implement all modules: authentication, "
     "hostel management, allocation, payments, and reports."),
    ("Phase 4 — Testing",
     "Conduct unit testing of individual views and models, integration testing of the "
     "allocation and payment workflows, and user acceptance testing (UAT) with selected "
     "stakeholders from the university."),
    ("Phase 5 — Deployment & Evaluation",
     "Deploy the system on a local server for demonstration. Evaluate against the original "
     "requirements using a structured checklist and gather feedback from stakeholders."),
]
for i, (title, body) in enumerate(phases, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r1 = p.add_run(title + "\n")
    r1.bold = True
    r1.font.color.rgb = NAVY
    r1.font.size = Pt(11)
    r2 = p.add_run(body)
    r2.font.size = Pt(11)

doc.add_paragraph()

# ── 7. SYSTEM FEATURES ────────────────────────────────────────────────────────

add_heading(doc, "7. KEY SYSTEM FEATURES", size=14)
section_rule(doc, "F0A500")

features = [
    ("Role-Based Access Control",
     "Two distinct roles — Administrator (Manager) and Student — each with a "
     "dedicated dashboard and access-controlled views."),
    ("Hostel/Block/Room Hierarchy",
     "Three-level physical structure: Hostel → Block → Room, with full CRUD "
     "management accessible to administrators through the web interface."),
    ("Smart Allocation",
     "Room allocation enforces capacity limits, prevents double-allocation of "
     "students, and validates gender-matching between student and hostel type."),
    ("Payment Management",
     "Students submit payment records with reference numbers; administrators verify "
     "or reject them, providing a clear audit trail."),
    ("Reports & Analytics",
     "Three report types — Allocation Report, Payment Report, Hostel Occupancy "
     "Report — all exportable to CSV for further analysis."),
    ("Responsive Design",
     "Bootstrap 5.3 ensures the system works on desktops, tablets, and mobile "
     "devices without a separate mobile application."),
    ("Security",
     "POST-only logout (CSRF protection), login-required guards on all private "
     "views, environment-based SECRET_KEY management, and role decorators on "
     "every sensitive endpoint."),
]
for title, body in features:
    p = doc.add_paragraph()
    r1 = p.add_run("• " + title + ": ")
    r1.bold = True
    r1.font.color.rgb = NAVY
    r1.font.size = Pt(11)
    r2 = p.add_run(body)
    r2.font.size = Pt(11)

doc.add_paragraph()

# ── 8. TECHNOLOGIES ───────────────────────────────────────────────────────────

add_heading(doc, "8. TECHNOLOGIES & TOOLS", size=14)
section_rule(doc, "F0A500")

tech_table = doc.add_table(rows=1, cols=3)
tech_table.style = 'Table Grid'
tech_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header row
hdr = tech_table.rows[0].cells
for cell, txt in zip(hdr, ["Category", "Technology / Tool", "Purpose"]):
    set_cell_bg(cell, "0D1F3C")
    r = cell.paragraphs[0].add_run(txt)
    r.bold = True
    r.font.color.rgb = WHITE
    r.font.size = Pt(10.5)

tech_data = [
    ("Backend Framework",  "Django 5.2 (Python 3.12)",      "MVC web framework, ORM, auth"),
    ("Database",           "SQLite 3",                       "Relational data storage"),
    ("Frontend",           "Bootstrap 5.3 + Bootstrap Icons","Responsive UI components"),
    ("Charts",             "Chart.js 4.4",                   "Dashboard doughnut charts"),
    ("Typography",         "Google Fonts — Inter",           "Modern, readable typeface"),
    ("Version Control",    "Git / GitHub",                   "Source code management"),
    ("IDE",                "VS Code",                        "Development environment"),
    ("Documentation",      "Microsoft Word / python-docx",  "Project documentation"),
]
for cat, tool, purpose in tech_data:
    row = tech_table.add_row().cells
    for cell, txt in zip(row, [cat, tool, purpose]):
        r = cell.paragraphs[0].add_run(txt)
        r.font.size = Pt(10.5)

doc.add_paragraph()

# ── 9. WORK PLAN ─────────────────────────────────────────────────────────────

add_heading(doc, "9. WORK PLAN / TIMELINE", size=14)
section_rule(doc, "F0A500")

plan_table = doc.add_table(rows=1, cols=4)
plan_table.style = 'Table Grid'
plan_table.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr = plan_table.rows[0].cells
for cell, txt in zip(hdr, ["Phase", "Activity", "Duration", "Period"]):
    set_cell_bg(cell, "0D1F3C")
    r = cell.paragraphs[0].add_run(txt)
    r.bold = True
    r.font.color.rgb = WHITE
    r.font.size = Pt(10.5)

plan_data = [
    ("1", "Requirements gathering & analysis",    "2 weeks",  "Week 1–2"),
    ("2", "System & database design",             "2 weeks",  "Week 3–4"),
    ("3", "Implementation — core modules",        "4 weeks",  "Week 5–8"),
    ("4", "Implementation — reports & UI polish", "2 weeks",  "Week 9–10"),
    ("5", "Testing & bug fixing",                 "2 weeks",  "Week 11–12"),
    ("6", "Documentation & proposal writing",     "1 week",   "Week 13"),
    ("7", "Presentation & evaluation",            "1 week",   "Week 14"),
]
for phase, activity, duration, period in plan_data:
    row = plan_table.add_row().cells
    for cell, txt in zip(row, [phase, activity, duration, period]):
        r = cell.paragraphs[0].add_run(txt)
        r.font.size = Pt(10.5)

doc.add_paragraph()

# ── 10. EXPECTED OUTCOMES ─────────────────────────────────────────────────────

add_heading(doc, "10. EXPECTED OUTCOMES", size=14)
section_rule(doc, "F0A500")

outcomes = [
    "A fully functional, web-based hostel management system accessible via any modern browser.",
    "A role-based system with distinct administrator and student interfaces.",
    "Reduction in manual paperwork and human error in room allocation and payment management.",
    "Real-time visibility of hostel occupancy for administrators.",
    "Exportable reports (CSV) for allocation and payment records.",
    "A secure, maintainable codebase following Django best practices.",
    "Complete project documentation including proposal, design diagrams, and user manual.",
]
for item in outcomes:
    add_bullet(doc, item)

doc.add_paragraph()

# ── 11. BUDGET (indicative) ───────────────────────────────────────────────────

add_heading(doc, "11. BUDGET ESTIMATE (Indicative)", size=14)
section_rule(doc, "F0A500")

budget_table = doc.add_table(rows=1, cols=3)
budget_table.style = 'Table Grid'
budget_table.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr = budget_table.rows[0].cells
for cell, txt in zip(hdr, ["Item", "Description", "Estimated Cost (UGX)"]):
    set_cell_bg(cell, "0D1F3C")
    r = cell.paragraphs[0].add_run(txt)
    r.bold = True
    r.font.color.rgb = WHITE
    r.font.size = Pt(10.5)

budget_data = [
    ("Internet / Data",       "Research, GitHub, CDN access",              "150,000"),
    ("Printing & Stationery", "Proposal, reports, forms",                  "80,000"),
    ("Transport",             "Site visits for requirements gathering",     "100,000"),
    ("Software",              "All tools are free/open-source",            "0"),
    ("Miscellaneous",         "Contingency",                               "70,000"),
    ("",                      "TOTAL",                                     "400,000"),
]
for item, desc, cost in budget_data:
    row = budget_table.add_row().cells
    for cell, txt in zip(row, [item, desc, cost]):
        r = cell.paragraphs[0].add_run(txt)
        r.font.size = Pt(10.5)
    if item == "":
        for cell in row:
            r = cell.paragraphs[0].runs[0]
            r.bold = True

doc.add_paragraph()

# ── 12. CONCLUSION ────────────────────────────────────────────────────────────

add_heading(doc, "12. CONCLUSION", size=14)
section_rule(doc, "F0A500")

add_body(doc,
    "The Online Hostel Information Management System addresses a genuine operational need at "
    "Cavendish University Uganda. By replacing manual, paper-based processes with a secure, "
    "role-aware web application, the system will improve accuracy, transparency, and efficiency "
    "in hostel administration. The project is technically feasible using industry-standard, "
    "open-source technologies and is achievable within the proposed timeline. The developer "
    "brings demonstrated competence in Django and web development, as evidenced by the working "
    "prototype already produced."
)
doc.add_paragraph()
add_body(doc,
    "It is therefore respectfully requested that this proposal be approved so that the project "
    "may proceed to full implementation, documentation, and evaluation as a final year project "
    "for the award of the Bachelor of Science in Information Technology at Cavendish University Uganda."
)

doc.add_paragraph()
doc.add_paragraph()

# ── SIGNATURE BLOCK ───────────────────────────────────────────────────────────

section_rule(doc, "F0A500")
doc.add_paragraph()

sig_table = doc.add_table(rows=4, cols=2)
sig_data = [
    ("Student's Signature:", "______________________________"),
    ("Name:",                "Natuhwera Marion"),
    ("Date:",                datetime.date.today().strftime("%B %d, %Y")),
    ("Supervisor's Approval:","______________________________"),
]
for i, (label, value) in enumerate(sig_data):
    row = sig_table.rows[i]
    r1 = row.cells[0].paragraphs[0].add_run(label)
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = row.cells[1].paragraphs[0].add_run(value)
    r2.font.size = Pt(11)

# ── SAVE ──────────────────────────────────────────────────────────────────────

output_path = (
    r"C:\Users\Marion\Desktop\HOSTEL MANAGEMENT SYSTEM\hostel_system"
    r"\Project_Proposal_HMS_Natuhwera_Marion.docx"
)
doc.save(output_path)
print(f"Document saved to:\n{output_path}")
